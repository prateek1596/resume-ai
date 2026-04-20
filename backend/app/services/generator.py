"""Generate resume HTML from structured data using Claude."""
from __future__ import annotations

from io import BytesIO
import html
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

try:
    import anthropic
except ImportError:  # pragma: no cover - depends on local environment
    anthropic = None

from app.core.config import get_settings
from app.models.resume import ResumeData

settings = get_settings()
client = anthropic.Anthropic(api_key=settings.anthropic_api_key) if anthropic else None

COLOR_SCHEMES = {
    "classic":  {"primary": "#1a1a2e", "accent": "#6c63ff", "text": "#2d2d3a", "light": "#f4f4f8", "mid": "#6b7280"},
    "navy":     {"primary": "#0f172a", "accent": "#3b82f6", "text": "#1e293b", "light": "#f0f9ff", "mid": "#64748b"},
    "emerald":  {"primary": "#064e3b", "accent": "#10b981", "text": "#1a2e20", "light": "#f0fdf4", "mid": "#6b7280"},
    "crimson":  {"primary": "#450a0a", "accent": "#dc2626", "text": "#2d1515", "light": "#fff5f5", "mid": "#9ca3af"},
    "slate":    {"primary": "#1e293b", "accent": "#475569", "text": "#334155", "light": "#f8fafc", "mid": "#94a3b8"},
    "gold":     {"primary": "#1c1a00", "accent": "#b45309", "text": "#1c1100", "light": "#fffbeb", "mid": "#92400e"},
}

TEMPLATE_DESCRIPTIONS = {
    "executive":  "Professional header with solid colored background, optional circular photo top-left, two-column skill tags, classic section dividers with accent color underlines.",
    "minimal":    "Ultra-clean white, single 3px left accent bar spanning full height, generous whitespace, small caps section labels, no colors except the single accent bar.",
    "split":      "Fixed 240px colored sidebar with contact+skills+photo, white main content area, sidebar uses primary color, content uses accent headings.",
    "classic":    "Traditional recruiter-friendly layout, thin ruled lines separating sections, bold section headers, two-line experience header (role+company / dates).",
    "creative":   "Colored sidebar with large circular photo, main content with decorative section markers, modern sans-serif typography.",
    "tech":       "Dark header band with monospace job title, code-like accent elements, developer-focused with tech stack badges.",
    "elegant":    "Centered name in large serif, gold/accent horizontal rule, photo centered below name, refined proportions.",
    "sharp":      "Angled polygon header, bold typography with high contrast, strong geometric accents.",
    "timeline":   "Vertical timeline dots connecting experience items, header with photo, chronological emphasis.",
    "ats_pure":   "Purely text-based, Arial only, no decorative elements, no colors, maximum ATS compatibility, all caps bold section headers.",
    "bento":      "Grid-based modular layout with stacked content cards, strong section rhythm, clean spacing, and photo in the top card.",
    "monograph":  "Editorial serif-led layout with centered masthead, generous margins, thin rules, and subdued magazine-style hierarchy.",
    "duo":        "Compact two-rail layout with a narrow side column for contact details and a dense main column for experience and projects.",
    "finance":    "Elegant finance-focused resume with precise metrics, restrained typography, and high-contrast section dividers.",
    "product":    "Outcome-led product resume with strategic summary placement, emphasis on impact metrics, and clean modular sections.",
    "portfolio":  "Portfolio-forward resume with stronger project showcase, visual hierarchy for links, and modern creative balance.",
}


def _looks_like_html(value: str) -> bool:
    return bool(re.search(r"<\s*(div|section|h1|h2|h3|p|ul|li|span|header|main)\b", value, re.IGNORECASE))


def _looks_like_latex(value: str) -> bool:
    return bool(re.search(r"\\documentclass|\\begin\{|\\end\{|\\section\{|\\resume", value))


def _strip_code_fences(value: str) -> str:
    return re.sub(r"^```[a-zA-Z]*\s*|```$", "", value.strip(), flags=re.MULTILINE).strip()


def _escape(value: str) -> str:
    return html.escape(value or "", quote=True)


def _render_fallback_resume(resume: ResumeData, color_scheme: str) -> str:
    cs = COLOR_SCHEMES.get(color_scheme, COLOR_SCHEMES["classic"])

    def section_title(text: str) -> str:
        return f'<div style="margin:22px 0 10px 0;font-size:12px;font-weight:700;letter-spacing:.12em;text-transform:uppercase;color:{cs["accent"]}">{_escape(text)}</div>'

    def tag(text: str) -> str:
        return f'<span style="display:inline-block;padding:5px 10px;border-radius:999px;background:{cs["light"]};color:{cs["text"]};font-size:12px;margin:0 6px 6px 0;">{_escape(text)}</span>'

    def optional_div(style: str, text: str) -> str:
        return f'<div style="{style}">{text}</div>' if text else ""

    contact_bits = [
        resume.contact.email,
        resume.contact.phone,
        resume.contact.location,
        resume.contact.linkedin,
        resume.contact.portfolio,
        resume.contact.website,
    ]
    contact_bits = [bit for bit in contact_bits if bit]

    experience_html = []
    for item in resume.experiences:
        bullets = "".join(f'<li style="margin:0 0 6px 18px;line-height:1.5">{_escape(b)}</li>' for b in item.bullets)
        end_text = f" - {_escape(item.end)}" if item.end else ""
        bullets_html = f'<ul style="margin:8px 0 0 0;padding:0;">{bullets}</ul>' if bullets else ""
        experience_html.append(
            f'<div style="margin-bottom:14px;page-break-inside:avoid">'
            f'<div style="display:flex;justify-content:space-between;gap:12px;font-weight:700">'
            f'<div>{_escape(item.role)}</div>'
            f'<div style="color:{cs["mid"]};font-weight:500;text-align:right">{_escape(item.start)}{end_text}</div>'
            f'</div>'
            f'<div style="margin-top:2px;color:{cs["text"]};font-weight:600">{_escape(item.company)}</div>'
            f'{bullets_html}'
            f'</div>'
        )

    education_html = []
    for item in resume.educations:
        degree_html = f'<div style="color:{cs["mid"]};font-size:13px">{_escape(item.degree)}</div>' if item.degree else ""
        year_html = f'<div style="color:{cs["mid"]};font-size:13px">{_escape(item.year)}</div>' if item.year else ""
        education_html.append(
            f'<div style="margin-bottom:10px">'
            f'<div style="font-weight:700">{_escape(item.school)}</div>'
            f'{degree_html}'
            f'{year_html}'
            f'</div>'
        )

    cert_html = []
    for item in resume.certifications:
        issuer_html = f'<div style="color:{cs["mid"]};font-size:13px">{_escape(item.issuer)}</div>' if item.issuer else ""
        cert_html.append(
            f'<div style="margin-bottom:8px">'
            f'<div style="font-weight:600">{_escape(item.name)}</div>'
            f'{issuer_html}'
            f'</div>'
        )

    project_html = []
    for item in resume.projects:
        bullets = "".join(f'<li style="margin:0 0 6px 18px;line-height:1.5">{_escape(b)}</li>' for b in item.bullets)
        technologies_html = f'<div style="color:{cs["mid"]};font-size:13px">{_escape(item.technologies)}</div>' if item.technologies else ""
        description_html = f'<div style="margin-top:4px;line-height:1.5">{_escape(item.description)}</div>' if item.description else ""
        bullets_html = f'<ul style="margin:8px 0 0 0;padding:0;">{bullets}</ul>' if bullets else ""
        project_html.append(
            f'<div style="margin-bottom:10px">'
            f'<div style="font-weight:700">{_escape(item.name)}</div>'
            f'{technologies_html}'
            f'{description_html}'
            f'{bullets_html}'
            f'</div>'
        )

    language_html = "".join(tag(f"{item.language}{f': {item.level}' if item.level else ''}") for item in resume.languages)
    skills_html = "".join(tag(skill) for skill in resume.skills.technical + resume.skills.soft)
    contact_line = " • ".join(_escape(bit) for bit in contact_bits)
    summary_html = optional_div("margin-top:12px;line-height:1.7;font-size:14px", _escape(resume.summary))
    title_html = optional_div(
        f"margin-top:4px;font-size:17px;color:{cs['mid']};font-weight:600",
        _escape(resume.contact.title),
    )
    contact_html = optional_div(
        f"margin-top:8px;font-size:12px;color:{cs['mid']};line-height:1.6",
        contact_line,
    )
    experience_section = section_title("Experience") if experience_html else ""
    education_section = section_title("Education") if education_html else ""
    cert_section = section_title("Certifications") if cert_html else ""
    skills_section = section_title("Skills") if skills_html else ""
    projects_section = section_title("Projects") if project_html else ""
    languages_section = section_title("Languages") if language_html else ""

    return (
        f'<div style="width:794px;min-height:1123px;padding:34px 40px;font-family:\'Segoe UI\',Arial,Helvetica,sans-serif;color:{cs["text"]};background:#fff;box-sizing:border-box">'
        f'<div style="border-bottom:3px solid {cs["accent"]};padding-bottom:14px;margin-bottom:8px">'
        f'<div style="font-size:34px;font-weight:800;line-height:1.1">{_escape(resume.contact.name or "Your Name")}</div>'
        f'{title_html}'
        f'{contact_html}'
        f'</div>'
        f'{summary_html}'
        f'{experience_section}'
        f'{"".join(experience_html)}'
        f'{education_section}'
        f'{"".join(education_html)}'
        f'{cert_section}'
        f'{"".join(cert_html)}'
        f'{skills_section}'
        f'{f"<div>{skills_html}</div>" if skills_html else ""}'
        f'{projects_section}'
        f'{"".join(project_html)}'
        f'{languages_section}'
        f'{f"<div>{language_html}</div>" if language_html else ""}'
        f'</div>'
    )

GENERATE_SYSTEM = """You are a senior front-end developer and professional resume designer. Generate a complete, beautiful, print-ready HTML resume.

Rules:
- Return ONLY the inner HTML content — a single <div> wrapping everything. No DOCTYPE, no <html>, no <head>, no <body>.
- Width exactly 794px (A4 width), use inline styles only — no <style> blocks, no external CSS.
- All fonts via inline font-family — use system fonts: 'Segoe UI', Arial, Helvetica, sans-serif (or Georgia for elegant).
- The resume must look exactly like a real professionally designed resume, not a generic one.
- Use the exact template style described. Do NOT default to a generic layout.
- Include ALL provided data — do not omit any section.
- Bullet points must start with strong action verbs.
- Photo: if photo_base64 is provided AND the template uses a photo, embed as <img src="{photo_base64}" style="width:80px;height:80px;object-fit:cover;border-radius:50%">
- ATS-friendly: use <h1>, <h2>, <h3>, semantic section structure, no tables for layout.
- Page breaks: avoid content being cut in the middle via page-break-inside:avoid on experience items."""


def _build_prompt(resume: ResumeData, template_id: str, color_scheme: str, job_description: str) -> str:
    cs = COLOR_SCHEMES.get(color_scheme, COLOR_SCHEMES["classic"])
    td = TEMPLATE_DESCRIPTIONS.get(template_id, TEMPLATE_DESCRIPTIONS["executive"])
    has_photo = bool(resume.photo_base64)

    contact_lines = [
        f"Name: {resume.contact.name}",
        f"Title: {resume.contact.title}",
        f"Email: {resume.contact.email}",
        f"Phone: {resume.contact.phone}",
        f"Location: {resume.contact.location}",
        f"LinkedIn: {resume.contact.linkedin}",
        f"Portfolio: {resume.contact.portfolio}",
        f"Website: {resume.contact.website}",
    ]
    exp_lines = []
    for e in resume.experiences:
        exp_lines.append(f"\n  {e.role} @ {e.company} ({e.start} — {e.end})")
        for b in e.bullets:
            exp_lines.append(f"    • {b}")

    edu_lines = [f"  {e.degree} in {e.field}, {e.school} ({e.year})" for e in resume.educations]
    cert_lines = [f"  {c.name} — {c.issuer} ({c.year})" for c in resume.certifications]
    lang_lines = [f"  {lang.language}: {lang.level}" for lang in resume.languages]
    proj_lines = []
    for p in resume.projects:
        proj_lines.append(f"\n  {p.name} | {p.technologies}")
        proj_lines.append(f"    {p.description}")

    photo_instruction = (
        f'Photo: EMBED this base64 image as <img src="{resume.photo_base64}" style="width:80px;height:80px;object-fit:cover;border-radius:50%"> in the appropriate template position.'
        if has_photo
        else "Photo: NONE — do not include any photo placeholder."
    )

    jd_note = f"\nTarget Job Description (optimize keywords):\n{job_description[:1500]}" if job_description else ""

    return f"""Generate a "{template_id}" template resume with this exact style:
{td}

Color Scheme:
- Primary (header bg / sidebar): {cs['primary']}
- Accent (section titles / links): {cs['accent']}
- Body text: {cs['text']}
- Light bg (skill pills / alternating): {cs['light']}

{photo_instruction}
{jd_note}

CANDIDATE DATA:
{chr(10).join(contact_lines)}

Summary:
{resume.summary}

Experience:{''.join(exp_lines) or ' None'}

Education:
{chr(10).join(edu_lines) or '  None'}

Technical Skills: {', '.join(resume.skills.technical) or 'None'}
Soft Skills: {', '.join(resume.skills.soft) or 'None'}

Certifications:
{chr(10).join(cert_lines) or '  None'}

Languages:
{chr(10).join(lang_lines) or '  None'}

Projects:{''.join(proj_lines) or ' None'}

Generate the complete resume HTML now:"""


async def generate(
    resume: ResumeData,
    template_id: str,
    color_scheme: str,
    job_description: str,
) -> str:
    if client is None:
        return _render_fallback_resume(resume, color_scheme)

    prompt = _build_prompt(resume, template_id, color_scheme, job_description)
    try:
        response = client.messages.create(
            model=settings.anthropic_model,
            max_tokens=4096,
            system=GENERATE_SYSTEM,
            messages=[{"role": "user", "content": prompt}],
        )

        from app.utils.html import sanitise
        raw = response.content[0].text
        if _looks_like_latex(raw):
            raise ValueError("Model returned LaTeX instead of HTML")
        # Strip any accidental markdown fences
        clean = re.sub(r"^```html\s*|^```\s*|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        if not _looks_like_html(clean):
            raise ValueError("Model returned non-HTML content")
        return sanitise(clean)
    except Exception:
        return _render_fallback_resume(resume, color_scheme)


async def improve_content(content: str, context: str, job_description: str, mode: str) -> tuple[str, list[str]]:
    """AI-improve a specific piece of resume content."""
    mode_instructions = {
        "bullets": "Rewrite these job bullets with strong action verbs, quantified achievements, and ATS-friendly language. Return one bullet per line, no bullet symbols.",
        "summary": "Rewrite this professional summary to be compelling, keyword-rich, and 2-3 sentences. Return only the improved text.",
        "general": "Improve this resume content for clarity, impact, and ATS optimization. Return only the improved text.",
    }

    instruction = mode_instructions.get(mode, mode_instructions["general"])
    jd_note = f"\nTarget role context:\n{job_description[:800]}" if job_description else ""

    if client is None:
        return content.strip(), [
            "Kept original content due to AI service unavailability",
            "Try again after configuring backend AI dependencies",
        ]

    try:
        response = client.messages.create(
            model=settings.anthropic_model,
            max_tokens=800,
            system=f"You are a professional resume writer. {instruction} Return ONLY the improved content, no commentary or explanation.",
            messages=[{"role": "user", "content": f"Context: {context}{jd_note}\n\nContent to improve:\n{content}"}],
        )
        improved = _strip_code_fences(response.content[0].text)
        if _looks_like_latex(improved):
            raise ValueError("Model returned LaTeX-like output")
        improved = improved or content.strip()
    except Exception:
        return content.strip(), [
            "Kept original content because AI improvement failed",
            "Check ANTHROPIC_API_KEY and model configuration",
        ]

    # Generate a brief list of what changed
    try:
        changes_response = client.messages.create(
            model=settings.anthropic_model,
            max_tokens=200,
            system="List 2-3 specific changes made. Return as JSON array of strings: [\"Added quantification\", \"Replaced weak verbs\"]. No markdown.",
            messages=[{"role": "user", "content": f"Original:\n{content}\n\nImproved:\n{improved}"}],
        )
        import json
        changes = json.loads(changes_response.content[0].text.strip())
        if not isinstance(changes, list) or not changes:
            raise ValueError("Invalid changes format")
    except Exception:
        changes = ["Improved clarity and impact", "Enhanced ATS compatibility"]

    return improved, changes


def build_docx_resume(resume: ResumeData) -> bytes:
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(resume.contact.name or "Your Name")
    title_run.bold = True
    title_run.font.size = Pt(18)

    if resume.contact.title:
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle.add_run(resume.contact.title).italic = True

    contact_bits = [
        resume.contact.email,
        resume.contact.phone,
        resume.contact.location,
        resume.contact.linkedin,
        resume.contact.portfolio,
        resume.contact.website,
    ]
    contact_line = " | ".join(bit for bit in contact_bits if bit)
    if contact_line:
        contact = document.add_paragraph()
        contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact.add_run(contact_line)

    if resume.summary:
        document.add_heading("Summary", level=1)
        document.add_paragraph(resume.summary)

    if resume.experiences:
        document.add_heading("Experience", level=1)
        for item in resume.experiences:
            heading = document.add_paragraph()
            heading.add_run(f"{item.role} | {item.company}").bold = True
            meta = document.add_paragraph()
            meta.add_run(f"{item.start} - {item.end}").italic = True
            for bullet in item.bullets:
                document.add_paragraph(bullet, style="List Bullet")

    if resume.projects:
        document.add_heading("Projects", level=1)
        for item in resume.projects:
            heading = document.add_paragraph()
            heading.add_run(item.name).bold = True
            if item.technologies:
                document.add_paragraph(item.technologies)
            if item.description:
                document.add_paragraph(item.description)
            for bullet in item.bullets:
                document.add_paragraph(bullet, style="List Bullet")

    if resume.educations:
        document.add_heading("Education", level=1)
        for item in resume.educations:
            line = f"{item.school} | {item.degree} {item.field}".strip()
            document.add_paragraph(line)
            if item.year:
                document.add_paragraph(item.year)

    if resume.skills.technical or resume.skills.soft:
        document.add_heading("Skills", level=1)
        if resume.skills.technical:
            document.add_paragraph("Technical: " + ", ".join(resume.skills.technical))
        if resume.skills.soft:
            document.add_paragraph("Soft: " + ", ".join(resume.skills.soft))

    if resume.certifications:
        document.add_heading("Certifications", level=1)
        for item in resume.certifications:
            line = item.name
            if item.issuer:
                line = f"{line} | {item.issuer}"
            if item.year:
                line = f"{line} | {item.year}"
            document.add_paragraph(line)

    if resume.languages:
        document.add_heading("Languages", level=1)
        for item in resume.languages:
            level = f": {item.level}" if item.level else ""
            document.add_paragraph(f"{item.language}{level}")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
